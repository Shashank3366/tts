from docx import Document

def create_sample_docx():
    doc = Document()
    doc.add_heading('TTS Test Sentences', 0)
    
    sentences = [
        "The quick brown fox jumps over the lazy dog.",
        "Artificial intelligence is transforming the world of voice synthesis.",
        "Hello, this is a test for 11Labs and Sarvam AI text to speech systems.",
        "Please ensure the audio quality is clear and natural.",
        "Testing short sentence.",
        "Testing a slightly longer sentence to see how the duration and pitch analysis handles it compared to the shorter ones."
    ]
    
    for sentence in sentences:
        doc.add_paragraph(sentence)
        
    doc.save('english_test_texts.docx')
    print("Created english_test_texts.docx")

if __name__ == "__main__":
    create_sample_docx()
